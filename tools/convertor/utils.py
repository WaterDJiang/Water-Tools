import streamlit as st
import pandas as pd
import io
import json

from fpdf import FPDF
import tempfile
from pathlib import Path
import os
import logging
import pdfplumber
from docx import Document

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.font_path = None
    def set_chinese_font(self):
        """设置中文字体，兼容多平台"""
        font_paths = [
            '/System/Library/Fonts/PingFang.ttc',
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/System/Library/Fonts/Hiragino Sans GB.ttc',
            os.path.join(os.path.dirname(__file__), 'fonts', 'SimSun.ttf'),
            'C:\\Windows\\Fonts\\SimSun.ttf',
            '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
        ]
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    self.font_path = font_path
                    self.add_font("CustomFont", fname=font_path)
                    return True
                except Exception as e:
                    logger.error(f"加载字体失败: {str(e)}")
                    continue
        logger.error("未找到可用的中文字体")
        return False
    def footer(self):
        self.set_y(-15)
        self.set_font('CustomFont', size=8)
        self.cell(0, 10, f'第 {self.page_no()} 页', align='C')

def json_to_pdf(content):
    """将JSON内容转换为PDF格式，保持原始格式，返回PDF文件路径或错误信息"""
    try:
        if isinstance(content, bytes):
            content = content.decode('utf-8')
        data = json.loads(content)
    except Exception as e:
        return False, f"JSON解析错误: {str(e)}"
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, 'output.pdf')
    try:
        pdf = PDF()
        pdf.set_margin(10)
        if not pdf.set_chinese_font():
            return False, "错误：未找到合适的中文字体，请确保已安装中文字体"
        pdf.add_page()
        pdf.set_font('CustomFont', size=9)
        pdf.set_auto_page_break(auto=True, margin=15)
        formatted_json = json.dumps(data, ensure_ascii=False, indent=2)
        for line in formatted_json.split('\n'):
            clean_line = ''.join(char for char in line if char.isprintable() or char in ['\n', '\t', ' '])
            if clean_line.strip():
                pdf.cell(0, 6, txt=clean_line, new_x="LMARGIN", new_y="NEXT")
        pdf.output(output_path)
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return True, output_path
        else:
            return False, "错误：生成的PDF文件为空"
    except Exception as e:
        return False, f"生成PDF时出错: {str(e)}"

def read_file(file, filetype):
    """根据文件类型读取数据，返回DataFrame或dict，异常时返回None和错误信息"""
    try:
        if filetype == "CSV":
            return pd.read_csv(file), None
        elif filetype == "Excel":
            return pd.read_excel(file), None
        elif filetype == "JSON":
            text = file.read()
            data = json.loads(text)
            if isinstance(data, list):
                return pd.DataFrame(data), None
            elif isinstance(data, dict):
                for v in data.values():
                    if isinstance(v, list):
                        return pd.DataFrame(v), None
                return pd.DataFrame([data]), None
            else:
                return None, "JSON结构不支持，仅支持对象或数组。"
        else:
            return None, "暂不支持的文件类型！"
    except Exception as e:
        return None, f"文件读取失败: {str(e)}"

def convert_and_download(df, target_format, filename_prefix="converted"):
    """根据目标格式转换并生成下载链接，异常时返回False和错误信息"""
    try:
        if target_format == "CSV":
            towrite = io.StringIO()
            df.to_csv(towrite, index=False)
            st.download_button(
                label="下载CSV文件",
                data=towrite.getvalue(),
                file_name=f"{filename_prefix}.csv",
                mime="text/csv"
            )
        elif target_format == "Excel":
            towrite = io.BytesIO()
            with pd.ExcelWriter(towrite, engine="openpyxl") as writer:
                df.to_excel(writer, index=False)
            st.download_button(
                label="下载Excel文件",
                data=towrite.getvalue(),
                file_name=f"{filename_prefix}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        elif target_format == "JSON":
            json_str = df.to_json(orient="records", force_ascii=False, indent=2)
            st.download_button(
                label="下载JSON文件",
                data=json_str,
                file_name=f"{filename_prefix}.json",
                mime="application/json"
            )
        else:
            st.error("暂不支持的目标格式！")
            return False, "暂不支持的目标格式！"
        return True, None
    except Exception as e:
        st.error(f"转换或下载失败: {str(e)}")
        return False, f"转换或下载失败: {str(e)}"

def pdf_to_docx(pdf_file):
    """PDF 转 Word，返回 docx 文件路径或错误信息"""
    try:
        if hasattr(pdf_file, 'read'):
            # 兼容 Streamlit UploadedFile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                tmp.write(pdf_file.read())
                pdf_path = tmp.name
        elif isinstance(pdf_file, str):
            pdf_path = pdf_file
        else:
            return False, "无法识别的 PDF 文件类型"
        doc = Document()
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        doc.add_paragraph(line)
        out_path = tempfile.mktemp(suffix='.docx')
        doc.save(out_path)
        if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
            return True, out_path
        else:
            return False, "生成的 Word 文件为空"
    except Exception as e:
        return False, f"PDF 转 Word 失败: {str(e)}"

def docx_to_pdf(docx_file):
    """Word 转 PDF，返回 PDF 文件路径或错误信息（仅文本，简单排版）"""
    try:
        if hasattr(docx_file, 'read'):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(docx_file.read())
                docx_path = tmp.name
        elif isinstance(docx_file, str):
            docx_path = docx_file
        else:
            return False, "无法识别的 Word 文件类型"
        doc = Document(docx_path)
        pdf = PDF()
        pdf.set_margin(10)
        if not pdf.set_chinese_font():
            return False, "错误：未找到合适的中文字体，请确保已安装中文字体"
        pdf.add_page()
        pdf.set_font('CustomFont', size=11)
        pdf.set_auto_page_break(auto=True, margin=15)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                pdf.cell(0, 8, txt=text, new_x="LMARGIN", new_y="NEXT")
        out_path = tempfile.mktemp(suffix='.pdf')
        pdf.output(out_path)
        if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
            return True, out_path
        else:
            return False, "生成的 PDF 文件为空"
    except Exception as e:
        return False, f"Word 转 PDF 失败: {str(e)}"

def pdf_tables_to_docx(pdf_file):
    """提取 PDF 中所有表格并导出为 Word 表格，返回 docx 文件路径或错误信息"""
    import pdfplumber
    from docx import Document
    try:
        if hasattr(pdf_file, 'read'):
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                tmp.write(pdf_file.read())
                pdf_path = tmp.name
        elif isinstance(pdf_file, str):
            pdf_path = pdf_file
        else:
            return False, "无法识别的 PDF 文件类型"
        doc = Document()
        table_count = 0
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        table_count += 1
                        word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                        for i, row in enumerate(table):
                            for j, cell in enumerate(row):
                                word_table.cell(i, j).text = str(cell) if cell is not None else ''
                        doc.add_paragraph()  # 表格间空行
        if table_count == 0:
            return False, "未检测到可提取的表格。"
        out_path = tempfile.mktemp(suffix='.docx')
        doc.save(out_path)
        if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
            return True, out_path
        else:
            return False, "生成的 Word 文件为空"
    except Exception as e:
        return False, f"PDF 表格提取失败: {str(e)}" 